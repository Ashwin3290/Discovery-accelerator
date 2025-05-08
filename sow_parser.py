import re
import os
from typing import Dict, List, Any, Tuple
import google.generativeai as genai
import docx
import fitz  # PyMuPDF for PDF processing
import io
from PIL import Image
import time
import traceback

class SOWParser:
    def __init__(self, gemini_api_key=None):
        # Set API key for Gemini
        if gemini_api_key:
            os.environ["GOOGLE_API_KEY"] = gemini_api_key
            genai.configure(api_key=gemini_api_key)
        elif "GOOGLE_API_KEY" in os.environ:
            genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
        else:
            raise ValueError("No Gemini API key provided. Please provide one or set GOOGLE_API_KEY environment variable.")
            
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        
    def parse_sow(self, file_path: str) -> Dict[str, Any]:
        print(f"\n=== STARTING SOW PARSING at {time.strftime('%Y-%m-%d %H:%M:%S')} ===")
        print(f"Parsing file: {file_path}")
        """
        Parse SOW document and extract sections, requirements, and boundaries
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            print(f"Detected PDF file, extracting text...")
            document_text = self._extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            print(f"Detected DOCX file, extracting text...")
            document_text = self._extract_text_from_docx(file_path)
        else:
            print(f"ERROR: Unsupported file format: {file_extension}")
            raise ValueError(f"Unsupported file format: {file_extension}")
            
        print(f"Successfully extracted {len(document_text)} characters of text")
        preview = document_text[:100].replace('\n', ' ').strip()
        print(f"Text preview: {preview}...")
        
        print("Parsing sections from Gemini response...")
        try:
            # Extract document structure using Gemini
            print(f"\nStep 1: Extracting document sections at {time.strftime('%Y-%m-%d %H:%M:%S')}...")
            sections = self._extract_sections(document_text)
            print(f"Successfully extracted {len(sections)} sections")
            
            print(f"\nStep 2: Extracting requirements at {time.strftime('%Y-%m-%d %H:%M:%S')}...")
            requirements = self._extract_requirements(sections)
            print(f"Successfully extracted {len(requirements)} requirements")
            
            print(f"\nStep 3: Identifying scope boundaries at {time.strftime('%Y-%m-%d %H:%M:%S')}...")
            boundaries = self._identify_boundaries(sections)
            print(f"Successfully identified scope boundaries: {len(boundaries.get('in_scope', []))} in-scope, " +
                  f"{len(boundaries.get('out_of_scope', []))} out-of-scope, {len(boundaries.get('unclear', []))} unclear")
        except Exception as e:
            print(f"ERROR during SOW parsing: {str(e)}")
            traceback.print_exc()
            raise
        
        print(f"\n=== COMPLETED SOW PARSING at {time.strftime('%Y-%m-%d %H:%M:%S')} ===\n")
        
        return {
            'sections': sections,
            'requirements': requirements,
            'boundaries': boundaries,
            'full_text': document_text
        }
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        print(f"Extracting text from PDF: {pdf_path}")
        """Extract text from PDF file"""
        text = ""
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            print(f"PDF has {page_count} pages")
            
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                text += page_text
                
                # Print progress for large documents
                if page_count > 5 and (page_num + 1) % 5 == 0:
                    print(f"Processed {page_num + 1}/{page_count} pages")
                    
            doc.close()
            print(f"Successfully extracted {len(text)} characters from PDF")
        except Exception as e:
            print(f"ERROR extracting text from PDF {pdf_path}: {str(e)}")
            traceback.print_exc()
        return text
    
    def _extract_text_from_docx(self, docx_path: str) -> str:
        print(f"Extracting text from DOCX: {docx_path}")
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(docx_path)
            print(f"DOCX has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract text from paragraphs
            para_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    para_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                table_count += 1
                
            print(f"Processed {para_count} non-empty paragraphs and {table_count} tables")
            print(f"Successfully extracted {len(text)} characters from DOCX")
        except Exception as e:
            print(f"ERROR extracting text from DOCX {docx_path}: {str(e)}")
            traceback.print_exc()
        return text
    
    def _extract_sections(self, document_text: str) -> Dict[str, str]:
        print(f"Extracting sections from document at {time.strftime('%Y-%m-%d %H:%M:%S')}...")
        """Extract sections from SOW document using Gemini"""
        prompt = f"""
        You are analyzing a Statement of Work (SOW) document. 
        Identify and extract the key sections from this document.
        Common sections in SOW documents include:
        - Introduction/Background
        - Scope of Work
        - Deliverables
        - Timeline/Schedule
        - Acceptance Criteria
        - Pricing/Payment Terms
        - Assumptions/Constraints
        - Change Management Process
        
        Extract each section with its heading and content.
        Format your response as a JSON object with section names as keys and their content as values.
        
        Document text:
        {document_text}
        """
        
        print(f"Sending section extraction prompt to Gemini API...")
        try:
            response = self.model.generate_content(prompt)
            print(f"Received response from Gemini API at {time.strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"Response length: {len(response.text)} characters")
        except Exception as api_error:
            print(f"ERROR: Gemini API call failed during section extraction: {str(api_error)}")
            traceback.print_exc()
            raise
        
        try:
            # Try to parse sections from the response
            sections_text = response.text
            # Extract JSON part from the response
            json_pattern = r'```json\n([\s\S]*?)\n```'
            json_match = re.search(json_pattern, sections_text)
            if json_match:
                import json
                sections = json.loads(json_match.group(1))
                print(f"Successfully parsed {len(sections)} sections from JSON response")
            else:
                print("JSON pattern not found in response, trying direct JSON parsing")
                # Try direct JSON parsing
                try:
                    import json
                    sections = json.loads(sections_text)
                    print(f"Successfully parsed {len(sections)} sections with direct JSON parsing")
                except:
                    print("Failed to parse JSON directly, falling back to heuristic approach")
                    # If no JSON found, use a simple heuristic approach
                    sections = {}
                    current_section = "Introduction"
                    lines = sections_text.split('\n')
                    section_content = []
                    
                    for line in lines:
                        if line.strip() and line[0] == '#' and len(line.strip()) < 50:
                            # If we find a new section heading, save the previous section
                            if section_content:
                                sections[current_section] = '\n'.join(section_content)
                            
                            # Start a new section
                            current_section = line.strip('#').strip()
                            section_content = []
                        else:
                            section_content.append(line)
                    
                    # Add the last section
                    if section_content:
                        sections[current_section] = '\n'.join(section_content)
            
            print(f"Sections extracted: {', '.join(list(sections.keys())[:5])}{'...' if len(sections) > 5 else ''}")
            return sections
                
        except Exception as e:
            print(f"ERROR extracting sections: {str(e)}")
            traceback.print_exc()
            # Fall back to a simple approach
            print("Falling back to using full document as a single section")
            return {"Full Document": document_text}
    
    def _extract_requirements(self, sections: Dict[str, str]) -> List[Dict[str, str]]:
        """Extract requirements from SOW sections using Gemini"""
        requirements = []
        
        # Combine relevant sections for requirement extraction
        relevant_sections = []
        for section_name, section_content in sections.items():
            if any(keyword in section_name.lower() for keyword in 
                  ['scope', 'deliverable', 'requirement', 'objective', 'service', 
                   'feature', 'function', 'specification', 'work', 'task', 'activity',
                   'responsibility', 'obligation', 'timeline', 'acceptance']):
                relevant_sections.append(f"## {section_name}\n{section_content}")
        
        if not relevant_sections:
            # If no specifically relevant sections, use all sections
            for section_name, section_content in sections.items():
                relevant_sections.append(f"## {section_name}\n{section_content}")
        
        combined_text = "\n\n".join(relevant_sections)
        
        prompt = f"""
        Extract all specific requirements from these SOW sections.
        Be comprehensive - consider only technical requirements and any statement that implies or connect to that to be done as a requirement.
        
        For each requirement:
        1. Provide a short ID (e.g., REQ-01)
        2. Extract the exact requirement text
        3. Identify which section it comes from
        4. Determine if it's clearly defined or ambiguous
        
        A requirement should be considered ambiguous if it:
        - Contains vague or subjective terms (e.g., "appropriate", "reasonable", "sufficient")
        - Lacks measurable criteria or specific details
        - Uses unclear terminology or jargon without definition
        - Has multiple possible interpretations
        - Doesn't specify who is responsible for the work
        - Contains conditional statements without clear triggers
        - Uses words like "may", "might", "could", "should" instead of "will", "shall", "must"
        - Doesn't have clear acceptance criteria
        - Lacks timeline or deadline information
        
        Format your response as a JSON array of objects with keys:
        - id: A unique identifier for the requirement
        - text: The exact requirement text
        - section: The section it comes from
        - clarity: Either "clear" or "ambiguous"
        - reason: Brief explanation if marked as ambiguous
        
        SOW Sections:
        {combined_text}
        """
        
        response = self.model.generate_content(prompt)
        print(f"Requirement extraction response length: {len(response.text)}")
        
        try:
            # Try to parse requirements from the response
            requirements_text = response.text
            # Extract JSON part from the response
            json_pattern = r'```json\n([\s\S]*?)\n```'
            json_match = re.search(json_pattern, requirements_text)
            if json_match:
                import json
                requirements = json.loads(json_match.group(1))
                print(f"Successfully parsed {len(requirements)} requirements from JSON response")
            else:
                # Try direct JSON parsing
                try:
                    import json
                    requirements = json.loads(requirements_text)
                    print(f"Successfully parsed {len(requirements)} requirements with direct JSON parsing")
                except:
                    print("Failed to parse JSON directly, attempting manual parsing")
                    # Simple parsing for requirements if no JSON found
                    lines = requirements_text.split('\n')
                    current_req = {}
                    
                    for line in lines:
                        if line.startswith('REQ-') or line.startswith('- ID: REQ-'):
                            # Save previous requirement if it exists
                            if current_req and 'id' in current_req:
                                requirements.append(current_req)
                            
                            # Start a new requirement
                            current_req = {'id': line.strip()}
                        elif ':' in line and current_req:
                            key, value = line.split(':', 1)
                            key = key.strip().lower()
                            if key in ['text', 'section', 'clarity', 'reason']:
                                current_req[key] = value.strip()
                    
                    # Add the last requirement
                    if current_req and 'id' in current_req:
                        requirements.append(current_req)
                    
                    print(f"Parsed {len(requirements)} requirements with manual parsing")
        
        except Exception as e:
            print(f"Error extracting requirements: {str(e)}")
            import traceback
            traceback.print_exc()
        
        # Make sure all requirements have the required fields
        for req in requirements:
            for field in ['id', 'text', 'section', 'clarity']:
                if field not in req:
                    req[field] = f"Unknown {field}"
            if req['clarity'] == 'ambiguous' and 'reason' not in req:
                req['reason'] = "No reason provided"
                
        print(f"Returning {len(requirements)} requirements")
        print(f"Ambiguous requirements: {sum(1 for req in requirements if req.get('clarity') == 'ambiguous')}")
        
        return requirements
    
    def _identify_boundaries(self, sections: Dict[str, str]) -> Dict[str, Any]:
        """Identify boundaries (in-scope vs out-of-scope) from SOW sections"""
        boundaries = {
            'in_scope': [],
            'out_of_scope': [],
            'unclear': []
        }
        
        # Combine relevant sections for boundary extraction
        scope_sections = []
        for section_name, section_content in sections.items():
            if any(keyword in section_name.lower() for keyword in 
                  ['scope', 'assumption', 'exclusion', 'limitation', 'constraint',
                   'boundary', 'deliverable', 'not included', 'included', 'exclude',
                   'include', 'work', 'service', 'responsibility']):
                scope_sections.append(f"## {section_name}\n{section_content}")
        
        if not scope_sections:
            # If no specifically relevant sections, use all sections
            for section_name, section_content in sections.items():
                scope_sections.append(f"## {section_name}\n{section_content}")
        
        combined_text = "\n\n".join(scope_sections)
        
        prompt = f"""
        Analyze these SOW sections and clearly identify what is in-scope, out-of-scope, and areas that are unclear or not explicitly defined.
        
        For the unclear/ambiguous areas, be liberal in your interpretation - any aspect that lacks specific detail or could be interpreted in multiple ways should be considered unclear.
        Consider the following as potentially unclear:
        - Services or deliverables without detailed specifications
        - Responsibilities without clear assignment
        - Processes without defined steps
        - Quality expectations without measurable criteria
        - Requirements with subjective terms (like "appropriate" or "reasonable")
        - Timeline dependencies without specific milestones
        - Resource requirements without quantification
        
        Format your response as a JSON object with three arrays:
        1. "in_scope": List of items explicitly included in the scope
        2. "out_of_scope": List of items explicitly excluded from the scope
        3. "unclear": List of important items that should be clarified (with brief explanation of why)
        
        SOW Sections:
        {combined_text}
        """
        
        response = self.model.generate_content(prompt)
        print(f"Boundary identification response length: {len(response.text)}")
        
        try:
            # Try to parse boundaries from the response
            boundaries_text = response.text
            # Extract JSON part from the response
            json_pattern = r'```json\n([\s\S]*?)\n```'
            json_match = re.search(json_pattern, boundaries_text)
            if json_match:
                import json
                boundaries = json.loads(json_match.group(1))
                print(f"Successfully parsed boundaries from JSON response")
            else:
                # Try direct JSON parsing
                try:
                    import json
                    boundaries = json.loads(boundaries_text)
                    print(f"Successfully parsed boundaries with direct JSON parsing")
                except:
                    print("Failed to parse JSON directly, attempting manual parsing")
                    # Simple parsing if no JSON found
                    in_scope_section = False
                    out_scope_section = False
                    unclear_section = False
                    
                    lines = boundaries_text.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        if "in-scope" in line.lower() or "in scope" in line.lower():
                            in_scope_section = True
                            out_scope_section = False
                            unclear_section = False
                            continue
                        elif "out-of-scope" in line.lower() or "out of scope" in line.lower():
                            in_scope_section = False
                            out_scope_section = True
                            unclear_section = False
                            continue
                        elif "unclear" in line.lower() or "ambiguous" in line.lower():
                            in_scope_section = False
                            out_scope_section = False
                            unclear_section = True
                            continue
                        
                        if line.startswith('-') or line.startswith('*'):
                            item = line[1:].strip()
                            if in_scope_section:
                                boundaries['in_scope'].append(item)
                            elif out_scope_section:
                                boundaries['out_of_scope'].append(item)
                            elif unclear_section:
                                boundaries['unclear'].append(item)
                    
                    print(f"Parsed boundaries with manual parsing")
        
        except Exception as e:
            print(f"Error identifying boundaries: {str(e)}")
            import traceback
            traceback.print_exc()
        
        # Make sure all key arrays exist
        for key in ['in_scope', 'out_of_scope', 'unclear']:
            if key not in boundaries:
                boundaries[key] = []
                
        print(f"Boundaries identified: {len(boundaries.get('in_scope', []))} in-scope, "
              f"{len(boundaries.get('out_of_scope', []))} out-of-scope, "
              f"{len(boundaries.get('unclear', []))} unclear")
        
        return boundaries
